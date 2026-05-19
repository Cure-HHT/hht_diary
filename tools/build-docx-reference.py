#!/usr/bin/env python3
"""Generate tools/urs-reference.docx from pandoc's default reference doc.

The reference doc supplies pandoc with the styles to use when emitting
the URS in Word format:

- Title (cover): centered, 24pt bold
- Heading 1 (chapters): bold, 18pt, page-break-before
- Heading 2 (sections): bold, 14pt
- Heading 3 (REQ headings): bold, 12pt
- Heading 4 (sub-sections under a REQ): bold italic, 11pt
- Body Text: 11pt Arial (Calibri fallback)
- Page header: sponsor / protocol / version, right-aligned
- Page footer: "Specific to Protocol TER-4480-C01    CONFIDENTIAL    Page N of M"

Run once (or whenever URS metadata changes); committed output lives at
tools/urs-reference.docx and is consumed automatically by
tools/compile-urs.py when producing docx output.

Usage:
  python3 tools/build-docx-reference.py
"""
from __future__ import annotations

import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


HEADER_TEXT = (
    "eCOA User Requirements Specification\n"
    "Sponsor: Terremoto Biosciences\n"
    "Protocol: TER-4480-C01    Version: 1.0"
)
FOOTER_LEFT = "Specific to Protocol TER-4480-C01"
FOOTER_CENTER = "CONFIDENTIAL"

# Body font preference: Arial (matches the URS LaTeX template), falling
# back to Calibri (Word default) if Arial isn't available on the renderer.
BODY_FONT = "Arial"


def _add_page_number_field(paragraph) -> None:
    """Insert a 'Page X of Y' field into the given paragraph."""
    run = paragraph.add_run("Page ")
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText_page = OxmlElement("w:instrText")
    instrText_page.set(qn("xml:space"), "preserve")
    instrText_page.text = " PAGE "
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText_page)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_end)

    paragraph.add_run(" of ")

    run2 = paragraph.add_run()
    fldChar_begin2 = OxmlElement("w:fldChar")
    fldChar_begin2.set(qn("w:fldCharType"), "begin")
    instrText_pages = OxmlElement("w:instrText")
    instrText_pages.set(qn("xml:space"), "preserve")
    instrText_pages.text = " NUMPAGES "
    fldChar_sep2 = OxmlElement("w:fldChar")
    fldChar_sep2.set(qn("w:fldCharType"), "separate")
    fldChar_end2 = OxmlElement("w:fldChar")
    fldChar_end2.set(qn("w:fldCharType"), "end")
    run2._r.append(fldChar_begin2)
    run2._r.append(instrText_pages)
    run2._r.append(fldChar_sep2)
    run2._r.append(fldChar_end2)


def _get_style(doc, style_name: str):
    """Look up a style by display name (workaround for python-docx 1.2.0
    where `doc.styles["Heading 1"]` raises KeyError; iterate instead)."""
    for s in doc.styles:
        if s.name == style_name:
            return s
    raise KeyError(f"no style with display name {style_name!r}")


def _set_heading_style(doc, style_name: str, size_pt: int, bold: bool = True,
                       italic: bool = False, page_break_before: bool = False,
                       color: tuple[int, int, int] = (0x1F, 0x3A, 0x5F)) -> None:
    """Configure a Heading N style for size, weight, color, and pagination."""
    style = _get_style(doc, style_name)
    font = style.font
    font.name = BODY_FONT
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(*color)
    pf = style.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    if page_break_before:
        pf.page_break_before = True


def _set_body_style(doc) -> None:
    style = _get_style(doc, "Normal")
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)


def _set_title_style(doc) -> None:
    # 'Title' is the style pandoc applies to the document's leading-most
    # paragraph when content carries a `# Title` or custom-style="Title".
    # We size and center it for the cover page experience.
    try:
        title = _get_style(doc, "Title")
    except KeyError:
        return
    title.font.name = BODY_FONT
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(24)


def _install_header(doc) -> None:
    """Replace each section's default header with the URS sponsor/protocol block."""
    for section in doc.sections:
        header = section.header
        # Clear existing paragraphs and rebuild with our content.
        for p in list(header.paragraphs):
            p._element.getparent().remove(p._element)
        for line_idx, line in enumerate(HEADER_TEXT.split("\n")):
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(line)
            run.font.name = BODY_FONT
            run.font.size = Pt(9)
            if line_idx == 0:
                run.bold = True


def _install_footer(doc) -> None:
    """Footer: 'Specific to Protocol ...    CONFIDENTIAL    Page X of Y'.

    python-docx doesn't expose tab stops on default footer paragraphs in
    a clean way; we use a 3-column table so the three labels land left,
    center, and right reliably.
    """
    for section in doc.sections:
        footer = section.footer
        # Clear any existing paragraphs
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)
        table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
        # Make the table look like plain text (no borders by default; nothing to do).
        # Left cell
        left = table.cell(0, 0).paragraphs[0]
        left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = left.add_run(FOOTER_LEFT)
        run.font.size = Pt(9)
        # Center cell
        center = table.cell(0, 1).paragraphs[0]
        center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = center.add_run(FOOTER_CENTER)
        run.font.size = Pt(9)
        run.bold = True
        # Right cell — page number field
        right = table.cell(0, 2).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Set the font on the runs we add
        _add_page_number_field(right)
        for r in right.runs:
            r.font.size = Pt(9)


def main() -> None:
    repo_root = Path(__file__).resolve().parent.parent
    output = repo_root / "tools" / "urs-reference.docx"

    # Start from pandoc's default reference doc — it ships with all the
    # style names pandoc expects to map markdown elements to.
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        default_path = Path(tmp.name)
    subprocess.run(
        ["pandoc", "--print-default-data-file", "reference.docx"],
        stdout=default_path.open("wb"),
        check=True,
    )

    doc = Document(default_path)

    _set_body_style(doc)
    _set_title_style(doc)
    _set_heading_style(doc, "Heading 1", size_pt=18, page_break_before=True)
    _set_heading_style(doc, "Heading 2", size_pt=14)
    _set_heading_style(doc, "Heading 3", size_pt=12)
    _set_heading_style(doc, "Heading 4", size_pt=11, italic=True)
    _set_heading_style(doc, "Heading 5", size_pt=11, bold=False, italic=True,
                       color=(0x40, 0x40, 0x40))

    _install_header(doc)
    _install_footer(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    main()
